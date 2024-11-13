import tkinter as tk
from tkinter import filedialog, ttk, simpledialog, Scrollbar, Text
import os
import docx
import zipfile
import threading
import base64
from io import BytesIO
import io
import PyPDF2
import re
import time
from datetime import datetime
import json
import subprocess

root_folder = ''
def browse_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        global root_folder
        root_folder = folder_path.replace("/", "\\")
        reset_tree()
        list_files(root_folder)
 
def reset_tree():
    for item in tree.get_children():
        tree.delete(item)
 
def list_files(folder_path, parent=""):
    try:
        total_files = len(os.listdir(folder_path))
    except Exception as e:
        tree.insert(parent, 'end', text="Folder doesn't exist")
        print(e)
        return
    progress_var.set(0)
    for i, item in enumerate(os.listdir(folder_path)):
        item_path = os.path.join(folder_path, item)
        if os.path.isdir(item_path):
            folder_name = os.path.basename(item_path)
            folder_id = tree.insert(parent, 'end', text=folder_name, open=True, image=image_folder)
            list_files(item_path, folder_id)
        elif item_path.endswith('.zip'):
            with zipfile.ZipFile(item_path, 'r') as zip_file:
                zip_folder_name = os.path.basename(item_path)
                zip_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=zip_folder_name, open=False, image=image_zip)
                for item in zip_file.namelist():
                    if not item.endswith('/'):
                        file_extension = os.path.splitext(item)[1].lower()
                        item_id = tree.insert(zip_id, 'end', text=item, image=get_image(file_extension))
        else:
            file_name = os.path.basename(item_path)
            file_extension = os.path.splitext(item)[1].lower()
            item_id = None
            item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, image=get_image(file_extension))
        
        if parent == "":
            progress = (i + 1) / total_files * 100
            progress_var.set(progress)
            app.update_idletasks()

def search_files_tag():
    global stop_work_flag
    stop_work_flag = False
    query = search_entry.get()
    if query == "":
        return
 
    reset_tree()
    # Create a new thread to perform a search task
    search_thread = threading.Thread(target=list_files_with_search_tag, args=(root_folder, query))
    search_thread.start()


def list_files_with_search_tag(folder_path, query, parent=""):
    global stop_work_flag
    total_files = len(os.listdir(folder_path))

    progress_var.set(0)  # Set the initial value of the progress bar
    for i, item in enumerate(os.listdir(folder_path)):
        # Check if the task should be stopped
        if stop_work_flag:
            break

        item_path = os.path.join(folder_path, item)
        file_name = os.path.basename(item_path)
        file_extension = os.path.splitext(item)[1].lower()
        if os.path.isdir(item_path):
            folder_name = os.path.basename(item_path)
            folder_id = tree.insert(parent, 'end', text=folder_name, open=True, image=image_folder)
            list_files_with_search_tag(item_path, query, folder_id)
        else:
            file_name = os.path.basename(item_path)
            get_tags_for_path(item_path, settings)
            if set(query.replace(", ", ",").split(",")).issubset(set(get_tags_for_path(item_path, settings).replace(", ", ",").split(","))):
                item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, tags=("green_check",), image=get_image(file_extension))
            else:
                item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, image=get_image(file_extension))

        # Update the progress bar
        progress = (i + 1) / total_files * 100
        progress_var.set(progress)
        app.update_idletasks()  # Force refresh the UI

def search_files():
    global stop_work_flag
    stop_work_flag = False
    query = search_entry.get()
    if query == "":
        return
 
    reset_tree()
    # Create a new thread to perform a search task
    search_thread = threading.Thread(target=list_files_with_search, args=(root_folder, query))
    search_thread.start()
 
def list_files_with_search(folder_path, query, parent=""):
    global stop_work_flag
    total_files = len(os.listdir(folder_path))
    # Compile the query string into a regular expression
    query_pattern = re.compile(query, re.IGNORECASE)  # Ignore case
    
    progress_var.set(0)  # Set the initial value of the progress bar
    for i, item in enumerate(os.listdir(folder_path)):
        # Check if the task should be stopped
        if stop_work_flag:
            break

        item_path = os.path.join(folder_path, item)
        file_name = os.path.basename(item_path)
        file_extension = os.path.splitext(item)[1].lower()
        if os.path.isdir(item_path):
            folder_name = os.path.basename(item_path)
            folder_id = tree.insert(parent, 'end', text=folder_name, open=True, image=image_folder)
            list_files_with_search(item_path, query, folder_id)
        elif item_path.endswith('.zip'):
            with zipfile.ZipFile(item_path, 'r') as zip_file:
                zip_folder_name = os.path.basename(item_path)
                zip_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=zip_folder_name, open=False, image=image_zip)
                for item in zip_file.namelist():
                    try:
                        if not item.endswith('/'):
                            file_extension = os.path.splitext(item)[1].lower()
                            if item.endswith(".docx"):
                                with zip_file.open(item) as word_content:
                                    doc = docx.Document(io.BytesIO(word_content.read()))
                                    content = ""
                                    for paragraph in doc.paragraphs:
                                        content += paragraph.text + "\n"
                                    if query_pattern.search(content):
                                        item_id = tree.insert(zip_id, 'end', text=item, tags=("green_check",), image=image_word)
                                    else:
                                        item_id = tree.insert(zip_id, 'end', text=item, image=image_word)
                            else:
                                with zip_file.open(item) as zip_content:
                                    file_content = zip_content.read()
                                    if query_pattern.search(file_content.decode('utf-8')):
                                        item_id = tree.insert(zip_id, 'end', text=item, tags=("green_check",), image=get_image(file_extension))
                                    else:
                                        item_id = tree.insert(zip_id, 'end', text=item, image=get_image(file_extension))
                    except Exception as e:
                        print("Fail to handle: " + item)
                        print(e)
                        item_id = tree.insert(zip_id, 'end', text=item, image=get_image(file_extension))
        else:
            file_name = os.path.basename(item_path)
            if query_pattern.search(file_name):
                item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, tags=("green_check",), image=get_image(file_extension))
            else:
                if file_extension == ".docx":
                    try:
                        doc = docx.Document(item_path)
                        content = ""
                        for paragraph in doc.paragraphs:
                            content += paragraph.text + "\n"
                        if query_pattern.search(content):
                            item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, tags=("green_check",), image=image_word)
                        else:
                            item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, image=image_word)
                    except:
                        item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, image=image_word)
                elif file_extension == ".pdf":
                    pdf_reader = PyPDF2.PdfReader(open(item_path, 'rb'))
                    content = ""
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        content += page.extract_text()
                    if query_pattern.search(content):
                        item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, tags=("green_check",), image=image_pdf)
                    else:
                        item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, image=image_pdf)
                else:
                    try:
                        with open(item_path, 'r', encoding='utf-8') as file:
                            file_content = file.read()
                            if query_pattern.search(file_content):
                                item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, tags=("green_check",), image=get_image(file_extension))
                            else:
                                item_id = tree.insert(parent, 'end', values=get_file_info(item_path), text=file_name, image=get_image(file_extension))
                    except Exception as e:
                        print("Fail to handle: " + item_path)
                        print(e)
                        item_id = tree.insert(parent, 'end', text=file_name, image=get_image(file_extension))
        # Update the progress bar
        progress = (i + 1) / total_files * 100
        progress_var.set(progress)
        app.update_idletasks()  # Force refresh the UI
 
def get_full_path(item):
    item_text = tree.item(item, "text")
    parent_item = tree.parent(item)
    if parent_item:
        parent_path = get_full_path(parent_item)
        return os.path.join(parent_path, item_text)
    else:
        return item_text

def open_file(event):
    try:
        selection = tree.selection()
        if selection:
            item = selection[0]
            item_text = tree.item(item, "text")
            if item_text:
                file_path = os.path.join(root_folder, get_full_path(item))
                os.startfile(file_path)
    except Exception as e:
        print(f"fail to open: {e}")

def set_default_path():
    try:
        global root_folder
        folder_path = filedialog.askdirectory()
        if folder_path:
            settings["default_path"] = folder_path.replace("/", "\\")
            with open('setting.json', 'w') as file:
                json.dump(settings, file, indent=4)
            root_folder = settings["default_path"]
            reset_tree()
            list_files(root_folder)
    except Exception as e:
        print(f"fail set default path: {e}")

def clear_search():
    global root_folder
    search_entry.delete(0, tk.END)
    reset_tree()
    list_files(root_folder)
    text_area.delete(1.0, tk.END)
 
# Increase the font size
def increase_font_size():
    current_font = text_area.cget("font")
    size = int(current_font.split(" ")[-1])
    size += 1
    new_font = ("Helvetica", size)
    text_area.configure(font=new_font)
    font_label.config(text=f"{size}")
 
# Decrease the font size
def decrease_font_size():
    current_font = text_area.cget("font")
    size = int(current_font.split(" ")[-1])
    if size > 1:
        size -= 1
        new_font = ("Helvetica", size)
        text_area.configure(font=new_font)
        font_label.config(text=f"{size}")

def tag(event):
    selection = tree.selection()
    if selection:
        item = selection[0]
        item_text = tree.item(item, "text")
        if item_text:
            file_path = os.path.join(root_folder, get_full_path(item))
            current_tags = get_tags_for_path(file_path, settings)
            new_tags = custom_askstring_with_tags("Update Tags", initialvalue=current_tags)
            set_tag_for_path(file_path, new_tags, settings)

def set_tag_for_path(file_path, new_tags, settings):
    not_in_setting = True
    for entry in settings["paths"]:
        if entry["path"] == file_path:
            entry["tag"] = new_tags
            not_in_setting = False
            break
    if not_in_setting:
        new_path = {
            "path": file_path,
            "tag": new_tags
        }
        settings["paths"].append(new_path)
    with open('setting.json', 'w') as file:
        json.dump(settings, file, indent=4)
    reset_tree()
    list_files(root_folder)

def custom_askstring_with_tags(title, initialvalue="", width=100):
    # Create a new window as a dialog
    dialog = tk.Toplevel()
    dialog.title(title)
    dialog.geometry("400x150")

    def add_initial_tag(tags):
        for tag in tags:
            tag_button = tk.Button(show_frame, text=tag, command=lambda t=tag: remove_tag(t))
            tag_button.pack(side="left", padx=5)
    def add_tag(tag):
        if tag != "":
            tags.append(tag)
            tag_button = tk.Button(show_frame, text=tag, command=lambda t=tag: remove_tag(t))
            tag_button.pack(side="left", padx=5)
            tag_entry.delete(0, "end")
    def remove_tag(tag):
        tags.remove(tag)
        for widget in show_frame.winfo_children():
            if widget.cget("text") == tag:
                widget.destroy()
    def on_ok():
        result = tag_entry.get()
        dialog.result = result
        dialog.destroy()

    # Add a tooltip label
    show_frame = tk.Frame(dialog)
    show_frame.pack(padx=10, pady=10)
    # Add the label section
    tag_frame = tk.Frame(dialog)
    tag_frame.pack(padx=10, pady=10)
    tags = []
    if initialvalue != "":
        tags = initialvalue.replace(", ", ",").split(",")
        add_initial_tag(tags)
    # New tag
    tag_entry = tk.Entry(tag_frame, width=15)
    tag_entry.pack(side="left", padx=5)
    # Set the focus
    tag_entry.focus_set()
    # add button
    add_tag_button = tk.Button(tag_frame, text="Add", command=lambda: add_tag(tag_entry.get()))
    add_tag_button.pack(side="left", padx=5)
    # ok button
    ok_button = tk.Button(dialog, text="OK", command=on_ok)
    ok_button.pack(padx=10, pady=10)
    # Wait for the window to close
    dialog.wait_window()
    # Return the value and label entered by the user
    return tags

def expand_all():
    expand_tree(tree)
 
def collapse_all():
    collapse_tree(tree)
 
def expand_tree(tree, item=""):
    children = tree.get_children(item)
    for child in children:
        tree.item(child, open=True)
        expand_tree(tree, child)
 
def collapse_tree(tree, item=""):
    children = tree.get_children(item)
    for child in children:
        tree.item(child, open=False)
        collapse_tree(tree, child)

def stop_work():
    global stop_work_flag
    stop_work_flag = True

def get_image(file_extension):
    if file_extension == ".docx":
        return image_word
    elif file_extension == ".xlsx":
        return image_excel
    elif file_extension == ".pptx":
        return image_ppt
    elif file_extension == ".pdf":
        return image_pdf
    elif file_extension == ".zip":
        return image_zip
    elif file_extension == ".mp4":
        return image_video
    elif file_extension in (".jpg", ".jpeg", ".png", ".gif"):
        return image_image
    elif file_extension in (".exe", "cmd"):
        return image_exe
    elif file_extension in (".dll", ".ini"):
        return image_setting
    elif file_extension in (".html", ".js", ".py", ".css", ".lang", ".cs", ".bat", ".css", ".cpp", ".pm"):
        return image_code
    else:
        return image_file

def format_size(size):
    kb = size / 1024
    mb = kb / 1024
    if mb >= 1:
        return f"{mb:.2f} MB"
    elif kb >= 1:
        return f"{kb:.2f} KB"
    else:
        return f"{size} bytes"

def get_file_info(file_path):
    file_size = os.path.getsize(file_path)
    file_size_str = format_size(file_size)
    modified_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y/%m/%d %H:%M:%S")
    tag = get_tags_for_path(file_path, settings)
    return (file_size_str, modified_time, tag)

def load_settings():
    settings = {"default_path": "", "paths": []}
    if os.path.exists("setting.json"):
        with open("setting.json", "r") as file:
            settings = json.load(file)
    return settings

def get_tags_for_path(path, settings):
    for entry in settings["paths"]:
        if entry["path"] == path:
            return ", ".join(entry["tag"])
    return ""

# Add a global variable to store the user's command location
command_prompt = "> "

# Execute the command in the callback function of the execute button
def execute_command(event):
    # Get the text from the current `text_area` region
    current_text = text_area.get("1.0", tk.END)

    # Get the content after the previous `>` character as the command
    command_start = current_text.rfind(command_prompt)
    if command_start != -1:
        command = current_text[command_start + len(command_prompt):].strip()
    else:
        command = ""

    if command:
        try:
            if root_folder:
                os.chdir(root_folder)

            # Execute the command and capture the output
            result = subprocess.run(command, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

            # Display the output of the command
            text_area.insert(tk.END, "\n")  # Add a newline before output
            text_area.insert(tk.END, result.stdout)
            text_area.insert(tk.END, result.stderr)

            # Add the command prompt again (on the same line)
            text_area.insert(tk.END, f"{command_prompt}", "prompt")  # Add prompt with tag
            text_area.mark_set(tk.INSERT, tk.END)  # Set the cursor at the end of the text

            # Move the scrollbar to the bottom to display the latest result and command prompt line
            text_area.see(tk.END)
        except Exception as e:
            text_area.insert(tk.END, str(e) + "\n")
            # Add the prompt after the error as well
            text_area.insert(tk.END, f"\n{command_prompt}", "prompt")
            text_area.mark_set(tk.INSERT, tk.END)
            text_area.see(tk.END)
    else:
        # If no command is entered, just display the new command prompt line on the same line
        # Only insert a prompt if the previous line has no input (avoiding an extra empty line)
        if not current_text.endswith(command_prompt + "\n"):
            text_area.insert(tk.END, f"{command_prompt}", "prompt")  # Add prompt with tag
        text_area.mark_set(tk.INSERT, tk.END)


app = tk.Tk()
app.title("LDMS")
settings = load_settings()
root_folder = settings["default_path"]
 
frame_top2 = tk.Frame(app)
frame_top2.grid(row=0, column=0, sticky="w", padx=10, pady=5)

setting_button = tk.Button(frame_top2, text="default", command=set_default_path)
setting_button.grid(row=1, column=0, padx=(0, 5), sticky="w")
 
browse_button = tk.Button(frame_top2, text="path", command=browse_folder)
browse_button.grid(row=1, column=1, padx=(0, 5))
 
expand_button = tk.Button(frame_top2, text="extend", command=expand_all)
expand_button.grid(row=1, column=2, padx=(0, 5), sticky="w")
 
collapse_button = tk.Button(frame_top2, text="close", command=collapse_all)
collapse_button.grid(row=1, column=3, padx=(0, 5), sticky="w")
 
frame_top = tk.Frame(app)
frame_top.grid(row=0, column=1, sticky="e", padx=10, pady=5)

# search
search_entry = tk.Entry(frame_top)
search_entry.grid(row=0, column=1, padx=(0, 5), ipady=10)
# search_entry.bind("<Return>", search_files)

search_button = tk.Button(frame_top, text="tag search", command=search_files_tag)
search_button.grid(row=0, column=2, padx=(0, 5))
 
search_button = tk.Button(frame_top, text="fuzzy search", command=search_files)
search_button.grid(row=0, column=3, padx=(0, 5))
 
clear_button = tk.Button(frame_top, text="clear", command=clear_search)
clear_button.grid(row=0, column=4, padx=(0, 5))
 
# "+" button
increase_font_button = tk.Button(frame_top, text="+", command=increase_font_size)
increase_font_button.grid(row=0, column=5, padx=(0, 5))
 
font_label = tk.Button(frame_top, text="12", command="")
font_label.grid(row=0, column=6, padx=(0, 5))
 
# "-" button
decrease_font_button = tk.Button(frame_top, text="-", command=decrease_font_size)
decrease_font_button.grid(row=0, column=7)
 
# Create a progress bar
progress_var = tk.DoubleVar()
progress_bar = ttk.Progressbar(app, variable=progress_var, maximum=100, mode='determinate')
progress_bar.grid(row=2, column=0, columnspan=2, padx=10, pady=(0, 5), sticky="nsew")
 
frame_bottom = ttk.PanedWindow(app, orient=tk.HORIZONTAL)
frame_bottom.grid(row=1, column=0, columnspan=2, padx=10, pady=5, sticky="nsew")
 
tree_frame = ttk.Frame(frame_bottom)
text_frame = ttk.Frame(frame_bottom)

# frame_bottom ratio
frame_bottom.add(tree_frame, weight=4)
frame_bottom.add(text_frame, weight=1)
frame_bottom.grid(row=1, column=0, columnspan=2, padx=10, pady=5, sticky="nsew")
frame_bottom.update_idletasks()

app.grid_rowconfigure(1, weight=1)  # Set the weight of row 1 to 1, making `frame_bottom` resizable
app.grid_columnconfigure(0, weight=1)  # Set the weight of column 0 to 1, making `frame_bottom` resizable

# tree_frame
tree = ttk.Treeview(tree_frame, columns=("Size", "Time", "Tag"), show="tree")
tree.column("Size", width=50)
tree.column("Time", width=100)
tree.column("Tag", width=100)
tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

vsb_tree = Scrollbar(tree_frame, orient="vertical", command=tree.yview)
vsb_tree.pack(side="right", fill="y")

tree.configure(yscrollcommand=vsb_tree.set)

# tree bottom
tree.bind("<Double-1>", open_file)
tree.bind("<Button-3>", tag)

# text_frame
text_area = Text(text_frame, wrap=tk.WORD, font=("Arial", 12), bg='black', fg='white')
text_area.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

# Set the tag for the prompt to make it green and uneditable
text_area.tag_configure("prompt", foreground="green", font=("Arial", 12, "bold"))
 
vsb = Scrollbar(text_frame, orient="vertical", command=text_area.yview)
vsb.pack(side="right", fill="y")
text_area.configure(yscrollcommand=vsb.set)

# Add the initial command prompt line
text_area.insert(tk.END, f"{command_prompt}")
text_area.bind("<Return>", execute_command)
 
# Adjust the row and column weights to ensure that `frame_bottom` resizes automatically
app.grid_rowconfigure(1, weight=1)
app.grid_columnconfigure(0, weight=1)
 
# Set the font size and spacing for the output area
style = ttk.Style()
style.configure("Treeview", font=("Arial", 13))
style.configure("Treeview.Heading", font=("Arial", 14))
 
# Set the appearance of the green checkmark at the beginning of the program.
tree.tag_configure("green_check", foreground="green", font=("Arial", 13, "bold"))
 
# Force stop
app.bind("<Escape>", lambda event: stop_work())

# Base64 encoding of images
icons = {
    "folder": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAABHNCSVQICAgIfAhkiAAAAAlwSFlzAAAAdgAAAHYBTnsmCAAAABl0RVh0U29mdHdhcmUAd3d3Lmlua3NjYXBlLm9yZ5vuPBoAAADdSURBVDiNY2AYaMD4YqFwLgMDwyRUUYbLbznemmiHMfwiZADTfwZGYQzR/wy6Ij9EcolxAQsuif////e8WCjcg0XqrUT8WxG4C4ixBQ2guBjFBVx8vAyMTIwETfh/WK4cwmDYDDeAlYOdgVdUkFhXdDAwMDAwMDLeh3uBk4+bWM0w8I7hN98mJohBjAwc3FykGrCU0fHADyYGBgYGDh5uBkYmEsOT6d98BgZoLJDsfEaGc4xWF88zMDAwMLFzcgizcrCTZsA/xnlwh3Bwc2KmRPzgAwPzz+Uk6qEhAAB6fSes95GcDgAAAABJRU5ErkJggg==",
    "file": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAABHNCSVQICAgIfAhkiAAAAAlwSFlzAAAAdgAAAHYBTnsmCAAAABl0RVh0U29mdHdhcmUAd3d3Lmlua3NjYXBlLm9yZ5vuPBoAAADhSURBVDiNndO/SgNBEMfxjxqwMTFqk9LX8BF8iYCF+BARrMTSwj5gI1j4HIIgloI2lv4JInqdohZ3S5Zwczn9wTY78/vu7DDDVAd4Dc47jrAg0Ape0A3il7jGGEt5YDEDTPARvYBhZb7A8ixgnr6whTNs4jQFOi0Bh9hR9uAO2+jjLSUMcNsSpsodRBWcYz0wXmGUX9QB9tELAI+zF3WAVeX/6lS0AexiLQDcKAeqEbAXmGvV1MRv5fA8/RWQmviJ5/9UcD/PlCuNcoEN8TLl6la5RV5BgRM8aFjZSj84ToBfH9spjYC2aR8AAAAASUVORK5CYII=",
    "word": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAABHNCSVQICAgIfAhkiAAAAAlwSFlzAAAAbwAAAG8B8aLcQwAAABl0RVh0U29mdHdhcmUAd3d3Lmlua3NjYXBlLm9yZ5vuPBoAAAE7SURBVDiNlZPNSkJRFIW/LTcS+7lSFqUQUWAFIXcslIEP0LhZEWT0CE3SV2gQYYQ+QeMGQVnTHqCROQg1DbQsMzNPIzXTq949O/us9e0F52xRSvG/RJCp3ZtFBENhm8ifrUc7RA3t/Pa1/X2YVRSGgIGyGaB8wFhDlIsGxAygfQxJUurMtlqdiV6Kb+Gu05GUBn/NJjHhqPuNQutnBji5K7adPbrGpm8UAJuZye/VB2G3AH6vju7QuDo0mJu0E99ftgY431tiJzDDisdBKOgmU6xaAzzmK4SCbuKJLFv+adKFL4uAXIVKtU7kIsVruUamMFiC5ivcJ0s8pMtUa3Vit1lKnz/WALFEptk8vnwayNwG6FUHa86egAx9fqOCSLe+IClt5Fst9Fsml3M8bAaXfussgus5unFqBvgFa+Vh8x7LAdEAAAAASUVORK5CYII=",
    "excel": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAABHNCSVQICAgIfAhkiAAAAAlwSFlzAAAAbwAAAG8B8aLcQwAAABl0RVh0U29mdHdhcmUAd3d3Lmlua3NjYXBlLm9yZ5vuPBoAAAFmSURBVDiNlZPNSgJRFMf/ZxxBQU3SEMVFTBlIKVOblkX5AC1ahArRGwQ9QNNsxFUIEiG17wks0icQjDIqSDBd2YdJfkCLstsixlJnGua/Opz743cPl3uIMYbhEBHNSgtTpk8mMqLxklzMjEAKOyktW2wfrTmO40QiJjLGRIDCAOwKdCUXSUvA23vtCnHkBRh+hhllG29tSfV2UJUnwKtl/wWxq37CYPIseVXtf2O2mPHUfYbgEpAv51FpViC4BAAAp0A+pw8BzzQAYMYTQDqW0vMOCpzWMRzE0wj7Q0iuJ5C9PjMmuK3fYT+XwvFWBoWHArKlU2MCAJiwudHsNhH0BWE1W40Jwv4QoosbiB9t4v6xDHlN4+GHwitF672N7ZMdvHQaSGSTWA2uGBPUXmv9Zu+rh/ObnDHBf4nOx/p1JBAZFDCgrvcbGbCn1idQle+YHILeMrmdDklLTnrrTAT3pXxxqCX4BuKVb+BooarYAAAAAElFTkSuQmCC",
    "ppt": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAABHNCSVQICAgIfAhkiAAAAAlwSFlzAAAAdgAAAHYBTnsmCAAAABl0RVh0U29mdHdhcmUAd3d3Lmlua3NjYXBlLm9yZ5vuPBoAAAFqSURBVDiNnZI9SJtRFIaf9yaiW9wKgkgiJGqh0KmQZildOphFyA+dxEWKW8diFwuu7eDgZINbDE4uKQ7i5M+g7dJWUD7B2cXJ0n73OOUzP19tmnc6cM77nJd7j/iLrDo+Bsk8pgLYZ9WDb3FzigyvsxnCsAD+OagATEd98x+RDnq3cJu0SmYbeEn4J9XF/Lfk3jlgDkg9PKgD1YMGjjNInKgeNFQPGsCV63PXI6tMVvFuGfxPK2cWWo1kq/BPX+CLi+jmGv04wn3ZbI/wAdiNPGIF2AC4TzA0jGvWSHxawqafdScYBXuMuTWwMmZvehI8KMMQMyjcB4XI5WIB/tU8lp9F3w87AWIdo4jkMHuvrfOLWIBr1nDHzZgE/hbpLRgIrJwuaivY6QT8/vU/FxApArjTvQHs7b8woAYDiBGrpEvAeBKpiucJWBaRA7LAcF8g86s972YlEgxlJwjDHLIpjBzer6tx+TWOcQfT0G7cnhISkwAAAABJRU5ErkJggg==",
    "pdf": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAMAAAAoLQ9TAAAAA3NCSVQICAjb4U/gAAAACXBIWXMAAABvAAAAbwHxotxDAAAAGXRFWHRTb2Z0d2FyZQB3d3cuaW5rc2NhcGUub3Jnm+48GgAAAHJQTFRF////8lVC8VZC4uXnwcbLsLe9vsTJytHY1tvf2cDB2re24uXn6K+p6p2U8VZC8mJP8mdV8mxa82pY83Bf9Hhp9Yx/9ZCC9pGE95uP956S96CV+Kmf+Kyi+bGo+bOq+bSr+beu+by0+r+3+sC5+8vF/NjT9D7P+AAAAAV0Uk5TAGCAuMCvAAwyAAAAZUlEQVQYV13OVxKAIAxF0SD22Hvv7n+LMkIs3L93ZsIAAMBRZoJKbbRMHUgIbBICdB3nD0I0QFdMZoRUgCgg+qRgXPK1j/f2gW0qm6I66weGLDnmtHtPtDeY4VPeDb9/SODv5nABrPoJ11tuehsAAAAASUVORK5CYII=",
    "zip": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAACXBIWXMAAAB2AAAAdgFOeyYIAAAAGXRFWHRTb2Z0d2FyZQB3d3cuaW5rc2NhcGUub3Jnm+48GgAAAglJREFUOI2Fkl1LVFEUhp/9MWdmHM+gNqZkDkQQZV0kSDf5E4oiuu4nBP2ELuqyLoIg/AsGQUERhBBBV6WV3tgHZoSijozzcZw5c85eXZyjzijYggWbtdd+9vuuvRWAvKSvzfCY8kYNaXgFF6jppVX+EwqgPcvnuHRjEv98z2aYn5gJS3feZTLkOx12AQcgws7IiHrbDVitTkfjNmM4HFEE1oJzSQIgTgY+mFvZ27yw3c0ioFS3vBCvtYgyCoPCRa09gNJ5/4nMnXjTA3AOTCpCuyqFrYeY5idwLciNQ7AMEicNQ2dOO+SulY+XLsTBjt+/9QBQ6NwIKEWmtcCXlX4ujw3A7k8QB8aHqNqlWFWsE31T9Q0O5OrPk2p7FDJDAPg5y6tvV7h27g/ENbDFHoBBfmhR6lfP1OLa/nJ+tczEqb/pwRqYAqiuQTv7XRthuRcQgEQAnB3eYKPhgymmdZfAUq1cnV/TeHb96NslKkp+g81aEXQWtJeoM/uAilKIJt+qHAGkNuZ/l5ksp5+xx4YFRR1Aq4tLIYAE20hzG2k3UrkdThbrbDb604nt2YjB+iAIgE7vDA7kh0inBVGNKNbkbSep62ySBzYCAJtQ3FTUNzjW7ULiXft6cWLGGAnvX5+9ByDKC4maKJv1TBwvHLF+OB4/e//10dO5YxvtcZs7a+tTKzT1cT3/AHCt1fBoXLBjAAAAAElFTkSuQmCC",
    "image": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAACXBIWXMAAAB2AAAAdgFOeyYIAAAAGXRFWHRTb2Z0d2FyZQB3d3cuaW5rc2NhcGUub3Jnm+48GgAAAUxJREFUOI3Fks8uA1EYxX/3zjShf5K2yTS0qURsKk2T7gix8BL2LLwBnsFb2IvEwo5YWCAsJCWCrlpRtBhTSpvO9Fqoqmq1FuK3vOc75zvJd+G/EQDzO6VBR9qrwDigdfE4wK5T1WZWpr03EqAm7UVgsgcz9ZkpTXcWACSAEiLSadrnEsT9OlK0CFJEAfRu65YSHgb6JRuXFdazlW+67KEyAKrDe9cGy8cloh6NM8vuHlBzHO5SacpmkWBsGG/Y4KmqOH18N9dsB6lr7QPKZpHM5h6v9xYADxcZwmMJQskYZdMiu33AS8HEGzYITyRxh4KfAfmj86HrwxOU3VRTKXL7Kcx0hrJVamjPuQLptS2C8ZHRRkAxexX5Ym7io1EzSikqpuWH+hWkS8+3df+A5nLdNgI8QWPObQQsROtvaYMQuA2/1ecLzP526d/wBg7dcjF0sLGyAAAAAElFTkSuQmCC",
    "video": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAACXBIWXMAAAB2AAAAdgFOeyYIAAAAGXRFWHRTb2Z0d2FyZQB3d3cuaW5rc2NhcGUub3Jnm+48GgAAAY9JREFUOI2tkjFrFFEUhb/7ZnY3+yZjllmXICgoGAsRQYOClXaCtYU2IZ2/JaWIv8FGbAQFU4idUUREELWIgiAW2d1idl4x+u5NNXGRXcmip3rcc8+5510u/CNkRi3Nsv5VgKoavlheXi3Mfp0HENEfk8n4A6BNs2senSODNZ/1H3e7xTETPWPYtvfFg8mkVZrZHcO21eS9z4pXS0srpw6mNUZJjI+Ac5Jw1qJck8QAue99IAS/4X0A4Saw7hL3EFgHSACybLCK2GmEN8CuEzseqvHddtvv4hi00roOYbzV6nQdwkdEvnfayU5d10EA8jzvq7YvLLI85+q3ZVkOU4AY08uIPVnEIMb0BvDUzWsQZEPgOvDur0nmEWa6V1WjZ6EaXRRk06BcyGAKqqrPBYZzDUTEZnG9Xq/Xzfpb4uQTcHKabDQpgHPyNeofHiL36p+yIlgxc7JzX+D3KTufFS+BS4f4EsBOqEZXAGt2oBrjLeD1YcQa423AphMcJMvzo2uqdmJ2bPlWlnufG/F/wT53Fov06YGj7AAAAABJRU5ErkJggg==",
    "exe": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAACXBIWXMAAAB2AAAAdgFOeyYIAAAAGXRFWHRTb2Z0d2FyZQB3d3cuaW5rc2NhcGUub3Jnm+48GgAAALlJREFUOI3F0jFqQkEQxvFfgn0KIdWTQCys02kTCy2sPYGVlYWd10iRA+QE3sHKNO8AFtrYpom1EAsXMsgKTxH8YGD2m9k/s+xwbz2EfIROxXtLfEWjj78LowePCfAaYD/YV5iiGQFRb3jHrgIkC5jhG4MqkBxggo8EGV4DuEi1jPeJKdqYY4FVqrXQzYHG/r/nxXEfftO5CH1F6Bufm6DEU6iV2Ka8ce4Jm+DVT3qeU5xqzQ1X+X46ANCZKIqF/7/AAAAAAElFTkSuQmCC",
    "code": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAABHNCSVQICAgIfAhkiAAAAAlwSFlzAAAAbwAAAG8B8aLcQwAAABl0RVh0U29mdHdhcmUAd3d3Lmlua3NjYXBlLm9yZ5vuPBoAAAF1SURBVDiNxZOxSyNBFMZ/b3YNuZNga4JnAmeyIFoJ1un8Nyyusb3uKhvB+z+u2EIRxEpRbFUQrDa7WIQghgWFK5Odze6zkCTkzq0EHXjFzDfv9z4+ZkRVec8y7+r+EID0O7/lKVgs1P/NQB6iGmoHurz+F0B64SOab2CyBCl90W9ev9CBhGGLodyQlFsAEkUVrJnXxmpMUq4zlCsJw9abAOl062TuGYn5pc3mNQDpnIc1EYA2m3ek5ifZ3Kl0uvUZgNz2qgz1nET2dW3lzwQ/wsNKNN7q6soh1uwx1HO57VWnDjJtY42iX6fNAInjkbjRzFle9rFGybQ9Aehmwyd1LxmMjiQISpPLqfFIpw4kCEoMRkek7qVuNvzZEC+WdkidZ+IFXw5wALDGI3nNQA5wiBd8UueZi6Wd/0LUXXKkto01OZV4S0Cw5jvu6B6ASryFNTlS29Zd8slgVX2zOOkvcxx3i/RxFb/EQfURy49CfZzLp//GF8XowBVVXkKkAAAAAElFTkSuQmCC",
    "setting": "iVBORw0KGgoAAAANSUhEUgAAABAAAAAQCAYAAAAf8/9hAAAABHNCSVQICAgIfAhkiAAAAAlwSFlzAAAAbwAAAG8B8aLcQwAAABl0RVh0U29mdHdhcmUAd3d3Lmlua3NjYXBlLm9yZ5vuPBoAAADdSURBVDiNndKxSgNREAXQM5JisQikEPIP6dMKEaxSpfB30gT8FLWUtFZ2qez8Aq2sUz+biTzWNXlxYGCZN/femdmrlKJOLLBHOZY//QPgr7phKAcJKvDiHILIgojY41JGKSU0RE1QmkFV70VD8yQilhHRDU43tFdVW+MTW3xgfuyIpQeeYYcpxrjDQ19k9MfYG7ziGk94xwuuft2rPwFWeEvl58xp1lYnV0ilG9wnuMMtNqeMdCB4zH3HqbzDrMWJB4J5Xnybf2DdauVSfXdYYnKOlf/lxFH/oYWgjm/fZ17fS/otRQAAAABJRU5ErkJggg=="
}
# Set icon
icon = base64.b64decode(icons.get("folder"))
image_folder = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("file"))
image_file = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("word"))
image_word = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("excel"))
image_excel = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("ppt"))
image_ppt = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("pdf"))
image_pdf = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("zip"))
image_zip = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("image"))
image_image = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("video"))
image_video = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("exe"))
image_exe = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("code"))
image_code = tk.PhotoImage(data=icon)
icon = base64.b64decode(icons.get("setting"))
image_setting = tk.PhotoImage(data=icon)
 
app.geometry("1200x900")
if root_folder != "":
    list_files(root_folder)
app.mainloop()
